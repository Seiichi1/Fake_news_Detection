import os
import re
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not found. Please install it using: pip install python-docx")
    exit(1)

# Configuration
MARKDOWN_FILE = r"C:\Users\rehan\.gemini\antigravity\brain\9bbd1cf1-317b-4b5b-b53a-923041d4a0a2\final_report.md.resolved"
OUTPUT_FILE = "Fake_News_Detection_Final_Report.docx"
DIAGRAMS_DIR = os.path.join(os.getcwd(), "diagrams")

# distinct mapping of placeholders to image files
IMAGE_MAPPING = {
    "app results": ["result1.png", "result2.png", "result3.png"],
    "loss curve": ["loss_diagram.png"],
    "accuracy curve": [], # No specific accuracy image found in list, might use classification_report or skip
    "confusion matrix": ["confusion_matrix.png"],
    "architecture": ["architecture_diagram.png"],
    "data strategy": ["data_flow_diagram.png"], # inferred
    "word clouds": ["word_cloud_real.png", "word_cloud_fake.png"],
}

def clean_text(text):
    """Removes markdown bold/italic markers for plain text insertion if needed, 
    but we will try to handle formatting."""
    text = text.strip()
    return text

def add_markdown_paragraph(doc, text):
    """Parses a paragraph for bold (**text**) and italic (*text*) and adds it to document."""
    if not text.strip():
        return
    
    p = doc.add_paragraph()
    
    # Simple parser: split by ** then *
    # This is a basic parser. For complex nested markdown it might need a real parser.
    # Current approach: Split by ** to separate bold parts.
    
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        bold = False
        if part.startswith('**') and part.endswith('**'):
            bold = True
            content = part[2:-2]
        else:
            content = part
            
        # Handle italics within the part
        subparts = re.split(r'(\*.*?\*)', content)
        for subpart in subparts:
            italic = False
            if subpart.startswith('*') and subpart.endswith('*') and len(subpart) > 1:
                italic = True
                subcontent = subpart[1:-1]
            else:
                subcontent = subpart
            
            if not subcontent:
                continue

            run = p.add_run(subcontent)
            run.bold = bold
            run.italic = italic

def add_table_from_lines(doc, lines):
    """Parses markdown table lines and adds a table to the doc."""
    # Filter header separator lines like |---|---|
    data_lines = [line for line in lines if not re.match(r'^\s*\|?[\s\-:|]+\|?\s*$', line)]
    
    if not data_lines:
        return

    # Determine dimensions
    rows = len(data_lines)
    cols = len(data_lines[0].strip().split('|')) - 2 # Assuming | content | format
    if cols < 1: cols = 1

    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'

    for r, line in enumerate(data_lines):
        # specific split to handle potential escaped pipes? (simplification: assume simple markdown tables)
        cells = line.strip().split('|')[1:-1]
        row_cells = table.rows[r].cells
        for c, cell_text in enumerate(cells):
            if c < len(row_cells):
                row_cells[c].text = cell_text.strip()

def insert_images_by_context(doc, context_text):
    """Checks context text to see if images should be inserted."""
    context_lower = context_text.lower()
    
    for key, filenames in IMAGE_MAPPING.items():
        if key in context_lower:
            for fname in filenames:
                path = os.path.join(DIAGRAMS_DIR, fname)
                if os.path.exists(path):
                    try:
                        doc.add_picture(path, width=Inches(6.0))
                        last_paragraph = doc.paragraphs[-1] 
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        print(f"Inserted image: {fname}")
                    except Exception as e:
                        print(f"Failed to insert {fname}: {e}")

def main():
    if not os.path.exists(MARKDOWN_FILE):
        print(f"Source file not found: {MARKDOWN_FILE}")
        return

    doc = Document()
    
    # Title
    doc.add_heading('Fake News Detection - Final Report', 0)

    with open(MARKDOWN_FILE, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    table_buffer = []
    in_table = False

    for line in lines:
        line = line.strip()
        
        # Skip title (already added)
        if line.startswith("# 🛡️"):
            continue

        # Tables
        if line.startswith("|"):
            table_buffer.append(line)
            in_table = True
            continue
        elif in_table:
            # End of table
            add_table_from_lines(doc, table_buffer)
            table_buffer = []
            in_table = False
        
        # Headers
        if line.startswith("## "):
            doc.add_heading(line.replace("## ", "").strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line.replace("### ", "").strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line.replace("#### ", "").strip(), level=3)
        
        # Images / Placeholders
        elif line.startswith("*(Please insert") or line.startswith("!["):
            # Try to deduce image from the line text
            insert_images_by_context(doc, line)
            
        # Horizontal Rule
        elif line == "---":
            doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Lists
        elif line.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            # remove dash
            content = line[2:]
            # Manual bold processing for list items
            # (Reusing add_markdown_paragraph logic but injecting into existing p is harder without refactoring)
            # Simplification: just add text to p run
            run = p.add_run(clean_text(content))
            # Basic bold check for whole line or parts?
            # Let's keep it simple for lists
            
        # Normal Text
        elif line:
            add_markdown_paragraph(doc, line)
            
            # Check if this paragraph context implies an image needs to follow
            # e.g. "Architecture Diagram:"
            # But the markdown places images explicitly usually. 
            # In the provided markdown, there are placeholders like "*(Please insert...)*" which are handled above.
            
    # Save
    doc.save(OUTPUT_FILE)
    print(f"Report generated successfully: {OUTPUT_FILE}")

if __name__ == "__main__":
    main()
